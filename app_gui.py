import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from fds_logic import process_fds_data
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg, NavigationToolbar2Tk
import re
import os
import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import datetime

import matplotlib
matplotlib.use('TkAgg')

# ─────────────────────────────────────────────────────────────
#  UI strings – Ukrainian / English
# ─────────────────────────────────────────────────────────────
TRANSLATIONS = {
    'uk': {
        'window_title': 'FDS Аналізатор Даних',
        'file_label': '1. Оберіть вхідний файл FDS (.txt або .csv):',
        'browse_btn': 'Огляд...',
        'column_mapping_btn': 'Налаштувати відповідність колонок',
        'bulk_mapping_btn': 'Масове призначення колонок',
        'run_btn': 'Запустити аналіз та відобразити графіки',
        'export_btn': 'Експорт даних у Word (.docx)',
        'status_waiting': 'Статус: Очікування вибору файлу...',
        'status_file_chosen': 'Статус: Файл вибрано: ',
        'status_cancelled': 'Статус: Вибір файлу скасовано.',
        'status_processing': 'Статус: Обробка даних та побудова графіків... Будь ласка, зачекайте.',
        'status_charts': 'Статус: Дані оброблені. Побудова графіків...',
        'status_done': 'Статус: Обробка завершена успішно! Графіки та зведена таблиця відображено.',
        'status_no_data': 'Статус: Обробка завершена, але дані відсутні.',
        'status_no_cols': 'Статус: Колонки не знайдено. Використайте налаштування відповідності.',
        'status_exporting': 'Статус: Експорт у Word... Будь ласка, зачекайте.',
        'status_export_done': 'Статус: Звіт Word успішно збережено: ',
        'err_no_file': 'Будь ласка, оберіть вхідний файл FDS (.txt або .csv).',
        'err_no_data': 'Дані для обробки не знайдено або вони порожні.',
        'err_no_cols': ('Не знайдено жодних колонок з даними для відображення.\n\n'
                        'Можливо, колонки у файлі мають нестандартні назви.\n'
                        "Спробуйте використати кнопку 'Налаштувати відповідність колонок'."),
        'err_no_export': 'Немає даних для експорту. Будь ласка, спочатку виконайте аналіз.',
        'err_read_file': 'Не вдалося прочитати файл: ',
        'err_export_word': 'Не вдалося зберегти файл Word: ',
        'err_file_not_found': 'Файл не знайдено: ',
        'err_unexpected': 'Виникла непередбачена помилка: ',
        'warn_no_file_mapping': 'Будь ласка, спочатку оберіть вхідний файл FDS.',
        'warn_no_data_cols': 'У файлі не знайдено колонок даних (крім часу).',
        'warn_no_numbered_cols': 'Не знайдено колонок з числовими індексами для масового призначення.',
        'warn_cancel_params': ('Ви відмовились вводити інформацію про нові параметри або ввели неповні дані. '
                               'Вони не будуть враховані.'),
        'info_mapping_done': 'Налаштовано відповідність для {n} колонок.',
        'info_bulk_done': 'Масово призначено {n} колонок.',
        'info_bulk_none': 'Не було вибрано жодної групи для призначення.',
        'info_export_done': 'Звіт Word успішно збережено за шляхом:\n',
        'info_success': 'Дані успішно оброблені, графіки та зведена таблиця відображено.',
        'info_cancelled_export': 'Експорт Word скасовано.',
        'dialog_mapping_title': 'Налаштування відповідності колонок',
        'dialog_mapping_header': 'Вкажіть, яка колонка відповідає якому параметру:',
        'dialog_mapping_hint': 'Залиште поле порожнім, якщо колонка не відповідає жодному параметру',
        'dialog_bulk_title': 'Масове призначення колонок',
        'dialog_bulk_header': 'Масове призначення параметрів для групи датчиків',
        'dialog_bulk_hint': 'Оберіть групу датчиків та призначте їм параметр',
        'dialog_unknown_params_title': 'Невідомі параметри',
        'dialog_unknown_params_header': 'Знайдено невідомі параметри. Будь ласка, вкажіть їх властивості:',
        'lbl_not_use': '-- Не використовувати --',
        'lbl_confirm': 'Підтвердити',
        'lbl_cancel': 'Скасувати',
        'lbl_count': 'Кількість: ',
        'lbl_examples': 'Приклади: ',
        'lbl_assign_param': 'Призначити параметр:',
        'lbl_show_cols': 'Показати всі колонки',
        'lbl_template': 'Шаблон: ',
        'lbl_param_code': 'Код: ',
        'lbl_full_name': 'Повна назва:',
        'lbl_units': 'Одиниці:',
        'lbl_critical_val': 'Критичне значення:',
        'lbl_direction': 'Напрямок:',
        'lbl_err_input': 'Помилка вводу',
        'lbl_all_fields': 'Будь ласка, заповніть всі поля для параметра ',
        'lbl_invalid_critical': 'Некоректне критичне значення для параметра ',
        'lbl_invalid_critical2': '. Введіть число.',
        'plot_sensor': 'Датчика',
        'plot_time_axis': 'Час (с)',
        'plot_critical_line': 'Критичний поріг',
        'plot_first_critical': 'Перший час досягнення критичного порогу: ',
        'plot_not_reached': 'Не досягнуто',
        'summary_title': 'Зведена таблиця критичних значень',
        'col_param': 'Параметр',
        'col_sensor': 'Сенсор',
        'col_critical': 'Критичне значення',
        'col_time': 'Час досягнення (с)',
        'word_report_title': 'Звіт аналізу даних FDS: ',
        'word_created': 'Дата та час створення звіту: ',
        'word_charts_heading': 'Графіки параметрів по датчиках',
        'word_table_heading': 'Зведена таблиця критичних значень',
        'word_chart_default': 'Графік',
        'file_dialog_title': 'Виберіть файл FDS (.txt або .csv)',
        'save_as_title': 'Зберегти звіт Word як',
        'param_Temp':  'Температура',
        'param_Visio': 'Видимість',
        'param_TP':    'Тепловий потік',
        'param_KK':    'Кисень',
        'param_OV':    'Оксид вуглецю',
        'param_DV':    'Діоксид вуглецю',
        'err_title':     'Помилка',
        'warn_title':    'Увага',
        'success_title': 'Успіх',
        'info_title':    'Інформація',
        'switch_lang_btn': 'EN',   # label shows the language you CAN switch TO
    },
    'en': {
        'window_title': 'FDS Data Analyzer',
        'file_label': '1. Select FDS input file (.txt or .csv):',
        'browse_btn': 'Browse...',
        'column_mapping_btn': 'Configure Column Mapping',
        'bulk_mapping_btn': 'Bulk Column Assignment',
        'run_btn': 'Run Analysis and Display Charts',
        'export_btn': 'Export Data to Word (.docx)',
        'status_waiting': 'Status: Waiting for file selection...',
        'status_file_chosen': 'Status: File selected: ',
        'status_cancelled': 'Status: File selection cancelled.',
        'status_processing': 'Status: Processing data and building charts... Please wait.',
        'status_charts': 'Status: Data processed. Building charts...',
        'status_done': 'Status: Processing complete! Charts and summary table displayed.',
        'status_no_data': 'Status: Processing complete, but no data found.',
        'status_no_cols': 'Status: No columns found. Use the column mapping settings.',
        'status_exporting': 'Status: Exporting to Word... Please wait.',
        'status_export_done': 'Status: Word report saved successfully: ',
        'err_no_file': 'Please select an FDS input file (.txt or .csv).',
        'err_no_data': 'No data found for processing or data is empty.',
        'err_no_cols': ('No data columns found to display.\n\n'
                        'The columns in the file may have non-standard names.\n'
                        'Try using the "Configure Column Mapping" button.'),
        'err_no_export': 'No data to export. Please run the analysis first.',
        'err_read_file': 'Failed to read file: ',
        'err_export_word': 'Failed to save Word file: ',
        'err_file_not_found': 'File not found: ',
        'err_unexpected': 'An unexpected error occurred: ',
        'warn_no_file_mapping': 'Please select an FDS input file first.',
        'warn_no_data_cols': 'No data columns found in the file (besides time).',
        'warn_no_numbered_cols': 'No columns with numeric indices found for bulk assignment.',
        'warn_cancel_params': ('You declined to provide information for new parameters or entered '
                               'incomplete data. They will not be included.'),
        'info_mapping_done': 'Mapping configured for {n} columns.',
        'info_bulk_done': 'Bulk assignment applied to {n} columns.',
        'info_bulk_none': 'No group was selected for assignment.',
        'info_export_done': 'Word report saved successfully to:\n',
        'info_success': 'Data processed successfully. Charts and summary table displayed.',
        'info_cancelled_export': 'Word export cancelled.',
        'dialog_mapping_title': 'Configure Column Mapping',
        'dialog_mapping_header': 'Specify which column corresponds to which parameter:',
        'dialog_mapping_hint': 'Leave blank if the column does not correspond to any parameter',
        'dialog_bulk_title': 'Bulk Column Assignment',
        'dialog_bulk_header': 'Bulk parameter assignment for sensor groups',
        'dialog_bulk_hint': 'Select a sensor group and assign a parameter to it',
        'dialog_unknown_params_title': 'Unknown Parameters',
        'dialog_unknown_params_header': 'Unknown parameters found. Please specify their properties:',
        'lbl_not_use': '-- Do not use --',
        'lbl_confirm': 'Confirm',
        'lbl_cancel': 'Cancel',
        'lbl_count': 'Count: ',
        'lbl_examples': 'Examples: ',
        'lbl_assign_param': 'Assign parameter:',
        'lbl_show_cols': 'Show all columns',
        'lbl_template': 'Template: ',
        'lbl_param_code': 'Code: ',
        'lbl_full_name': 'Full name:',
        'lbl_units': 'Units:',
        'lbl_critical_val': 'Critical value:',
        'lbl_direction': 'Direction:',
        'lbl_err_input': 'Input Error',
        'lbl_all_fields': 'Please fill in all fields for parameter ',
        'lbl_invalid_critical': 'Invalid critical value for parameter ',
        'lbl_invalid_critical2': '. Enter a number.',
        'plot_sensor': 'Sensor',
        'plot_time_axis': 'Time (s)',
        'plot_critical_line': 'Critical threshold',
        'plot_first_critical': 'First time critical threshold reached: ',
        'plot_not_reached': 'Not reached',
        'summary_title': 'Summary Table of Critical Values',
        'col_param': 'Parameter',
        'col_sensor': 'Sensor',
        'col_critical': 'Critical Value',
        'col_time': 'Time Reached (s)',
        'word_report_title': 'FDS Data Analysis Report: ',
        'word_created': 'Report created at: ',
        'word_charts_heading': 'Parameter Charts by Sensor',
        'word_table_heading': 'Summary Table of Critical Values',
        'word_chart_default': 'Chart',
        'file_dialog_title': 'Select FDS File (.txt or .csv)',
        'save_as_title': 'Save Word Report As',
        'param_Temp':  'Temperature',
        'param_Visio': 'Visibility',
        'param_TP':    'Heat Flux',
        'param_KK':    'Oxygen',
        'param_OV':    'Carbon Monoxide',
        'param_DV':    'Carbon Dioxide',
        'err_title':     'Error',
        'warn_title':    'Warning',
        'success_title': 'Success',
        'info_title':    'Information',
        'switch_lang_btn': 'UA',   # label shows the language you CAN switch TO
    },
}


class FDSAnalyzerApp:
    def __init__(self, master):
        self.master = master
        self.lang = 'uk'    # default language

        master.title(self.t('window_title'))
        master.geometry("1200x800")

        self.input_file_path = None
        self.processed_df = None
        self.processed_critical_points_data = None
        self.processed_parameters_info = None
        self.processed_parameter_order = None
        self.generated_figs = []
        self.summary_table_data = []

        # Language-independent base parameter info (no 'name' key – resolved via t())
        self.base_parameters_info = {
            'Temp':  {'unit': 'C',      'critical': 60.0,   'direction': 'above'},
            'Visio': {'unit': 'm',      'critical': 20.0,   'direction': 'below'},
            'TP':    {'unit': 'kW/m2',  'critical': 20.0,   'direction': 'above'},
            'KK':    {'unit': 'kg/m3',  'critical': 0.15,   'direction': 'below'},
            'OV':    {'unit': 'kg/m3',  'critical': 0.015,  'direction': 'above'},
            'DV':    {'unit': 'kg/m3',  'critical': 0.05,   'direction': 'above'},
        }
        self.current_parameters_info = self._build_parameters_info()
        self.column_mapping = {}

        self.unique_units = sorted(set(p['unit'] for p in self.base_parameters_info.values()))
        self.unique_critical_values = sorted(
            set(str(p['critical']) for p in self.base_parameters_info.values()))

        # ── Top bar (language toggle) ────────────────────────────────────
        top_bar = tk.Frame(master)
        top_bar.pack(fill="x", padx=20, pady=(10, 0))
        self.btn_lang = tk.Button(
            top_bar, text=self.t('switch_lang_btn'),
            command=self.toggle_language,
            font=("Arial", 10, "bold"), bg="#555555", fg="white", width=4)
        self.btn_lang.pack(side="right")

        # ── Main frame ──────────────────────────────────────────────────
        main_frame = tk.Frame(master, padx=20, pady=10)
        main_frame.pack(expand=True, fill="both")

        # 1. File selection
        self.lbl_file = tk.Label(main_frame, text=self.t('file_label'),
                                 font=("Arial", 10, "bold"))
        self.lbl_file.pack(anchor="w", pady=(0, 5))

        file_frame = tk.Frame(main_frame)
        file_frame.pack(fill="x", pady=(0, 10))
        self.entry_input = tk.Entry(file_frame, width=50)
        self.entry_input.pack(side="left", expand=True, fill="x")
        self.btn_browse_input = tk.Button(
            file_frame, text=self.t('browse_btn'), command=self.browse_input_file)
        self.btn_browse_input.pack(side="right")

        self.btn_column_mapping = tk.Button(
            main_frame, text=self.t('column_mapping_btn'),
            command=self.open_column_mapping_dialog,
            font=("Arial", 10, "bold"), bg="#FFA500", fg="white")
        self.btn_column_mapping.pack(pady=5)

        self.btn_bulk_mapping = tk.Button(
            main_frame, text=self.t('bulk_mapping_btn'),
            command=self.open_bulk_column_mapping_dialog,
            font=("Arial", 10, "bold"), bg="#FF6347", fg="white")
        self.btn_bulk_mapping.pack(pady=5)

        self.btn_run = tk.Button(
            main_frame, text=self.t('run_btn'),
            command=self.run_processing,
            font=("Arial", 12, "bold"), bg="#4CAF50", fg="white")
        self.btn_run.pack(pady=10)

        self.btn_export_doc = tk.Button(
            main_frame, text=self.t('export_btn'),
            command=self.export_results_to_doc,
            font=("Arial", 10, "bold"), bg="#1E90FF", fg="white",
            state=tk.DISABLED)
        self.btn_export_doc.pack(pady=5)

        self.status_label = tk.Label(
            main_frame, text=self.t('status_waiting'),
            fg="blue", font=("Arial", 10))
        self.status_label.pack(pady=10)

        # Scrollable charts area
        self.charts_container_frame = tk.Frame(main_frame, bd=2, relief="groove")
        self.charts_container_frame.pack(expand=True, fill="both", pady=10)

        self.canvas = tk.Canvas(self.charts_container_frame)
        self.canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        self.scrollbar = tk.Scrollbar(
            self.charts_container_frame, orient=tk.VERTICAL, command=self.canvas.yview)
        self.scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        self.canvas.configure(yscrollcommand=self.scrollbar.set)
        self.canvas.bind('<Configure>',
                         lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all")))

        self.charts_frame = tk.Frame(self.canvas)
        self.canvas.create_window((0, 0), window=self.charts_frame,
                                  anchor="nw", width=self.canvas.winfo_width())
        self.charts_frame.bind('<Configure>',
                               lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all")))
        self.canvas.bind('<Configure>', self.on_canvas_configure)

        self.canvas_widgets = []
        self.summary_table_widget = None

    # ── i18n helpers ────────────────────────────────────────────────────

    def t(self, key):
        """Return the translated string for *key* in the current language."""
        return TRANSLATIONS[self.lang].get(key, key)

    def _build_parameters_info(self):
        """Build parameters_info dict with names localised to current language."""
        result = {}
        for code, info in self.base_parameters_info.items():
            result[code] = {
                'name': self.t(f'param_{code}'),
                'unit': info['unit'],
                'critical': info['critical'],
                'direction': info['direction'],
            }
        return result

    def get_param_name(self, code):
        """Return localised name for a known parameter code (with fallback)."""
        key = f'param_{code}'
        if key in TRANSLATIONS[self.lang]:
            return TRANSLATIONS[self.lang][key]
        if code in self.current_parameters_info:
            return self.current_parameters_info[code].get('name', code)
        return code

    def toggle_language(self):
        """Switch between Ukrainian and English."""
        self.lang = 'en' if self.lang == 'uk' else 'uk'
        self.refresh_language()

    def refresh_language(self):
        """Update all static UI widget texts after a language change."""
        self.master.title(self.t('window_title'))
        self.lbl_file.config(text=self.t('file_label'))
        self.btn_browse_input.config(text=self.t('browse_btn'))
        self.btn_column_mapping.config(text=self.t('column_mapping_btn'))
        self.btn_bulk_mapping.config(text=self.t('bulk_mapping_btn'))
        self.btn_run.config(text=self.t('run_btn'))
        self.btn_export_doc.config(text=self.t('export_btn'))
        self.btn_lang.config(text=self.t('switch_lang_btn'))
        # Refresh known-param names in current_parameters_info
        for code in self.base_parameters_info:
            if code in self.current_parameters_info:
                self.current_parameters_info[code]['name'] = self.t(f'param_{code}')

    # ── Canvas helpers ───────────────────────────────────────────────────

    def on_canvas_configure(self, event):
        items = self.canvas.find_withtag("all")
        if items:
            self.canvas.itemconfig(items[0], width=event.width)

    # ── Column mapping dialogs ───────────────────────────────────────────

    def open_column_mapping_dialog(self):
        if not self.input_file_path:
            messagebox.showwarning(self.t('warn_title'), self.t('warn_no_file_mapping'))
            return
        try:
            df = pd.read_csv(self.input_file_path, header=[0, 1], skipinitialspace=True)
            new_columns = []
            for unit_header, param_sensor_header in df.columns:
                cu = unit_header.replace('"', '').replace("'", '')
                cp = param_sensor_header.replace('"', '').replace("'", '')
                new_columns.append(cu if 'Unnamed' in param_sensor_header else f"{cp} ({cu})")
            df.columns = new_columns
            if 's' in df.columns:
                df.rename(columns={'s': 'Time'}, inplace=True)
            elif 'Time (s)' in df.columns:
                df.rename(columns={'Time (s)': 'Time'}, inplace=True)

            available_columns = [c for c in df.columns if c != 'Time']
            if not available_columns:
                messagebox.showwarning(self.t('warn_title'), self.t('warn_no_data_cols'))
                return

            dialog = tk.Toplevel(self.master)
            dialog.title(self.t('dialog_mapping_title'))
            dialog.geometry("800x600")
            dialog.transient(self.master)
            dialog.grab_set()

            tk.Label(dialog, text=self.t('dialog_mapping_header'),
                     font=("Arial", 11, "bold")).pack(pady=10)
            tk.Label(dialog, text=self.t('dialog_mapping_hint'),
                     font=("Arial", 9), fg="gray").pack(pady=(0, 10))

            cf = tk.Frame(dialog)
            cf.pack(fill="both", expand=True, padx=10, pady=10)
            c = tk.Canvas(cf)
            sb = tk.Scrollbar(cf, orient="vertical", command=c.yview)
            sf = tk.Frame(c)
            sf.bind("<Configure>", lambda e: c.configure(scrollregion=c.bbox("all")))
            c.create_window((0, 0), window=sf, anchor="nw")
            c.configure(yscrollcommand=sb.set)
            c.pack(side="left", fill="both", expand=True)
            sb.pack(side="right", fill="y")

            mapping_vars = {}
            param_options = [self.t('lbl_not_use')] + [
                f"{code}: {self.get_param_name(code)}" for code in self.base_parameters_info]

            for col in available_columns:
                frame = tk.Frame(sf, pady=5)
                frame.pack(fill="x", padx=10)
                tk.Label(frame, text=col, width=40, anchor="w").pack(side="left", padx=5)
                tk.Label(frame, text="→", font=("Arial", 12)).pack(side="left", padx=5)
                var = tk.StringVar()
                if col in self.column_mapping:
                    pc = self.column_mapping[col]
                    var.set(f"{pc}: {self.get_param_name(pc)}" if pc in self.base_parameters_info
                            else self.t('lbl_not_use'))
                else:
                    var.set(self.t('lbl_not_use'))
                ttk.Combobox(frame, textvariable=var, values=param_options,
                             width=30, state="readonly").pack(side="left", padx=5)
                mapping_vars[col] = var

            bf = tk.Frame(dialog)
            bf.pack(pady=10)

            def on_confirm():
                new_mapping = {}
                for col, var in mapping_vars.items():
                    sel = var.get()
                    if sel != self.t('lbl_not_use'):
                        new_mapping[col] = sel.split(":")[0].strip()
                self.column_mapping = new_mapping
                messagebox.showinfo(self.t('success_title'),
                                    self.t('info_mapping_done').format(n=len(new_mapping)))
                dialog.destroy()

            def on_cancel():
                dialog.destroy()

            tk.Button(bf, text=self.t('lbl_confirm'), command=on_confirm,
                      bg="#4CAF50", fg="white", font=("Arial", 10, "bold"),
                      width=15).pack(side="left", padx=10)
            tk.Button(bf, text=self.t('lbl_cancel'), command=on_cancel,
                      bg="#f44336", fg="white", font=("Arial", 10, "bold"),
                      width=15).pack(side="left", padx=10)

        except Exception as e:
            messagebox.showerror(self.t('err_title'), self.t('err_read_file') + str(e))

    def open_bulk_column_mapping_dialog(self):
        if not self.input_file_path:
            messagebox.showwarning(self.t('warn_title'), self.t('warn_no_file_mapping'))
            return
        try:
            df = pd.read_csv(self.input_file_path, header=[0, 1], skipinitialspace=True)
            new_columns = []
            for unit_header, param_sensor_header in df.columns:
                cu = unit_header.replace('"', '').replace("'", '')
                cp = param_sensor_header.replace('"', '').replace("'", '')
                new_columns.append(cu if 'Unnamed' in param_sensor_header else f"{cp} ({cu})")
            df.columns = new_columns
            if 's' in df.columns:
                df.rename(columns={'s': 'Time'}, inplace=True)
            elif 'Time (s)' in df.columns:
                df.rename(columns={'Time (s)': 'Time'}, inplace=True)

            available_columns = [c for c in df.columns if c != 'Time']
            if not available_columns:
                messagebox.showwarning(self.t('warn_title'), self.t('warn_no_data_cols'))
                return

            pattern = re.compile(r'(.*?)\s*(\d+)\s*(\(.*\))$')
            grouped_columns = {}
            for col in available_columns:
                m = pattern.match(col)
                if m:
                    template = f"{m.group(1).strip()} {m.group(3)}"
                    grouped_columns.setdefault(template, []).append((col, m.group(2)))

            if not grouped_columns:
                messagebox.showinfo(self.t('info_title'), self.t('warn_no_numbered_cols'))
                return

            dialog = tk.Toplevel(self.master)
            dialog.title(self.t('dialog_bulk_title'))
            dialog.geometry("900x600")
            dialog.transient(self.master)
            dialog.grab_set()

            tk.Label(dialog, text=self.t('dialog_bulk_header'),
                     font=("Arial", 12, "bold")).pack(pady=10)
            tk.Label(dialog, text=self.t('dialog_bulk_hint'),
                     font=("Arial", 9), fg="gray").pack(pady=(0, 10))

            cf = tk.Frame(dialog)
            cf.pack(fill="both", expand=True, padx=10, pady=10)
            c = tk.Canvas(cf)
            sb = tk.Scrollbar(cf, orient="vertical", command=c.yview)
            sf = tk.Frame(c)
            sf.bind("<Configure>", lambda e: c.configure(scrollregion=c.bbox("all")))
            c.create_window((0, 0), window=sf, anchor="nw")
            c.configure(yscrollcommand=sb.set)
            c.pack(side="left", fill="both", expand=True)
            sb.pack(side="right", fill="y")

            mapping_vars = {}
            param_options = [self.t('lbl_not_use')] + [
                f"{code}: {self.get_param_name(code)}" for code in self.base_parameters_info]

            for template, columns_list in sorted(grouped_columns.items()):
                frame = tk.Frame(sf, pady=10, relief="ridge", bd=2)
                frame.pack(fill="x", padx=10, pady=5)

                hf = tk.Frame(frame, bg="#E8F4F8")
                hf.pack(fill="x", padx=5, pady=5)
                tk.Label(hf, text=self.t('lbl_template') + template,
                         font=("Arial", 10, "bold"), bg="#E8F4F8").pack(side="left", padx=5)
                tk.Label(hf, text=self.t('lbl_count') + str(len(columns_list)),
                         font=("Arial", 9), fg="gray", bg="#E8F4F8").pack(side="left", padx=5)
                examples = ", ".join(col for col, _ in columns_list[:3])
                if len(columns_list) > 3:
                    examples += "..."
                tk.Label(hf, text=self.t('lbl_examples') + examples,
                         font=("Arial", 8), fg="darkgray", bg="#E8F4F8").pack(side="left", padx=5)

                pf = tk.Frame(frame)
                pf.pack(fill="x", padx=10, pady=5)
                tk.Label(pf, text=self.t('lbl_assign_param'),
                         width=20, anchor="w").pack(side="left", padx=5)
                var = tk.StringVar(value=self.t('lbl_not_use'))
                ttk.Combobox(pf, textvariable=var, values=param_options,
                             width=35, state="readonly").pack(side="left", padx=5)

                def show_columns(cols=columns_list, tmpl=template):
                    messagebox.showinfo(f"'{tmpl}'",
                                        "\n".join(f"{col} (#{num})" for col, num in cols))

                tk.Button(pf, text=self.t('lbl_show_cols'),
                          command=show_columns, font=("Arial", 8)).pack(side="left", padx=5)
                mapping_vars[template] = (var, columns_list)

            bf = tk.Frame(dialog)
            bf.pack(pady=10)

            def on_confirm():
                count = 0
                for template, (var, columns_list) in mapping_vars.items():
                    sel = var.get()
                    if sel != self.t('lbl_not_use'):
                        pc = sel.split(":")[0].strip()
                        for col, num in columns_list:
                            self.column_mapping[col] = pc
                            count += 1
                if count > 0:
                    messagebox.showinfo(self.t('success_title'),
                                        self.t('info_bulk_done').format(n=count))
                else:
                    messagebox.showinfo(self.t('info_title'), self.t('info_bulk_none'))
                dialog.destroy()

            def on_cancel():
                dialog.destroy()

            tk.Button(bf, text=self.t('lbl_confirm'), command=on_confirm,
                      bg="#4CAF50", fg="white", font=("Arial", 10, "bold"),
                      width=15).pack(side="left", padx=10)
            tk.Button(bf, text=self.t('lbl_cancel'), command=on_cancel,
                      bg="#f44336", fg="white", font=("Arial", 10, "bold"),
                      width=15).pack(side="left", padx=10)

        except Exception as e:
            messagebox.showerror(self.t('err_title'), self.t('err_read_file') + str(e))

    # ── File browsing ────────────────────────────────────────────────────

    def browse_input_file(self):
        file_path = filedialog.askopenfilename(
            title=self.t('file_dialog_title'),
            filetypes=[("Data files", "*.txt *.csv"), ("All files", "*.*")]
        )
        if file_path:
            self.input_file_path = file_path
            self.entry_input.delete(0, tk.END)
            self.entry_input.insert(0, file_path)
            self.column_mapping = {}
            self.status_label.config(
                text=self.t('status_file_chosen') + os.path.basename(file_path), fg="green")
        else:
            self.status_label.config(text=self.t('status_cancelled'), fg="orange")

    # ── Chart helpers ────────────────────────────────────────────────────

    def clear_charts(self):
        for widget in self.canvas_widgets:
            widget.destroy()
        self.canvas_widgets.clear()
        for fig in self.generated_figs:
            plt.close(fig)
        self.generated_figs = []
        self.summary_table_data = []
        self.canvas.config(scrollregion=self.canvas.bbox("all"))

    # ── Unknown-parameter prompt ─────────────────────────────────────────

    def prompt_for_new_parameters(self, unknown_param_codes):
        dialog = tk.Toplevel(self.master)
        dialog.title(self.t('dialog_unknown_params_title'))
        dialog.transient(self.master)
        dialog.grab_set()

        tk.Label(dialog, text=self.t('dialog_unknown_params_header')).pack(pady=10)

        entries = {}
        for code in unknown_param_codes:
            frame = tk.Frame(dialog)
            frame.pack(pady=5, fill="x", padx=10)
            tk.Label(frame, text=self.t('lbl_param_code') + code).pack(side="left", padx=5)
            tk.Label(frame, text=self.t('lbl_full_name')).pack(side="left", padx=5)
            name_entry = tk.Entry(frame)
            name_entry.pack(side="left", expand=True, fill="x", padx=5)
            tk.Label(frame, text=self.t('lbl_units')).pack(side="left", padx=5)
            unit_cb = ttk.Combobox(frame, values=self.unique_units, width=8)
            unit_cb.pack(side="left", padx=5)
            unit_cb.set("")
            tk.Label(frame, text=self.t('lbl_critical_val')).pack(side="left", padx=5)
            crit_cb = ttk.Combobox(frame, values=self.unique_critical_values, width=8)
            crit_cb.pack(side="left", padx=5)
            crit_cb.set("")
            tk.Label(frame, text=self.t('lbl_direction')).pack(side="left", padx=5)
            dir_var = tk.StringVar(value="above")
            tk.Radiobutton(frame, text=">", variable=dir_var, value="above").pack(side="left")
            tk.Radiobutton(frame, text="<", variable=dir_var, value="below").pack(side="left")
            entries[code] = {'name': name_entry, 'unit': unit_cb,
                             'critical': crit_cb, 'direction': dir_var}

        def on_confirm():
            new_params = {}
            for code, w in entries.items():
                try:
                    name = w['name'].get()
                    unit = w['unit'].get()
                    crit_str = w['critical'].get()
                    critical = float(crit_str)
                    direction = w['direction'].get()
                    if name and unit and crit_str:
                        new_params[code] = {'name': name, 'unit': unit,
                                            'critical': critical, 'direction': direction}
                    else:
                        messagebox.showwarning(self.t('lbl_err_input'),
                                               self.t('lbl_all_fields') + code)
                        return
                except ValueError:
                    messagebox.showwarning(self.t('lbl_err_input'),
                                           self.t('lbl_invalid_critical') + code +
                                           self.t('lbl_invalid_critical2'))
                    return
            dialog.user_response = new_params
            dialog.destroy()

        tk.Button(dialog, text=self.t('lbl_confirm'), command=on_confirm).pack(pady=10)
        self.master.wait_window(dialog)
        return getattr(dialog, 'user_response', {})

    # ── Main processing ──────────────────────────────────────────────────

    def run_processing(self):
        if not self.input_file_path:
            messagebox.showwarning(self.t('warn_title'), self.t('err_no_file'))
            return

        self.status_label.config(text=self.t('status_processing'), fg="blue")
        self.master.update_idletasks()
        self.clear_charts()
        self.current_parameters_info = self._build_parameters_info()

        try:
            unknown_param_codes = []
            if self.column_mapping:
                mapped_codes = set(self.column_mapping.values())
                unknown_param_codes = [c for c in mapped_codes
                                       if c not in self.current_parameters_info]
            else:
                temp_df, _, _, _, _ = process_fds_data(
                    self.input_file_path, column_mapping=self.column_mapping)
                col_parser = re.compile(r'([A-Za-z]+)(\d+)\s*\((.*?)\)')
                file_codes = {col_parser.match(c).group(1)
                              for c in temp_df.columns if col_parser.match(c)}
                unknown_param_codes = [c for c in file_codes
                                       if c not in self.current_parameters_info]

            if unknown_param_codes:
                user_params = self.prompt_for_new_parameters(sorted(unknown_param_codes))
                if user_params:
                    self.current_parameters_info.update(user_params)
                else:
                    messagebox.showwarning(self.t('warn_title'), self.t('warn_cancel_params'))

            df, critical_points_data, parameters_info, parameter_order, plot_column_map = \
                process_fds_data(self.input_file_path,
                                 parameters_info_override=self.current_parameters_info,
                                 column_mapping=self.column_mapping)

            if df.empty:
                messagebox.showinfo(self.t('info_title'), self.t('err_no_data'))
                self.status_label.config(text=self.t('status_no_data'), fg="red")
                self.btn_export_doc.config(state=tk.DISABLED)
                return

            self.processed_df = df
            self.processed_critical_points_data = critical_points_data
            self.processed_parameters_info = parameters_info
            self.processed_parameter_order = parameter_order

            self.status_label.config(text=self.t('status_charts'), fg="blue")
            self.master.update_idletasks()

            print(f"DEBUG: plot_column_map has {len(plot_column_map)} elements")

            if len(plot_column_map) == 0:
                messagebox.showwarning(self.t('warn_title'), self.t('err_no_cols'))
                self.status_label.config(text=self.t('status_no_cols'), fg="orange")
                self.btn_export_doc.config(state=tk.DISABLED)
                return

            # Build per-sensor display items
            sensor_data = {}
            for (param_code, sensor_num), col_name in plot_column_map.items():
                sensor_data.setdefault(sensor_num, [])
                info = parameters_info[param_code]
                param_name = self.get_param_name(param_code)
                unit = info['unit']
                critical_threshold = info['critical']
                critical_key = f"{param_code}_{sensor_num}"

                critical_time_text = ""
                critical_time_for_summary = self.t('plot_not_reached')
                if critical_key in critical_points_data and critical_points_data[critical_key]:
                    crit_time = critical_points_data[critical_key][0]['time']
                    critical_time_text = self.t('plot_first_critical') + f"{crit_time:.2f} s"
                    critical_time_for_summary = f"{crit_time:.2f}"

                sensor_data[sensor_num].append({
                    "param_code": param_code,
                    "param_name": param_name,
                    "sensor_num": sensor_num,
                    "unit": unit,
                    "critical_threshold": critical_threshold,
                    "plot_col_name": col_name,
                    "critical_time_text": critical_time_text,
                    "critical_time_for_summary": critical_time_for_summary,
                    "param_order_idx": (list(parameters_info.keys()).index(param_code)
                                        if param_code in parameters_info
                                        else len(parameters_info)),
                })

            sorted_display_data = []
            for s_num in sorted(sensor_data.keys()):
                sorted_display_data.extend(
                    sorted(sensor_data[s_num], key=lambda x: x['param_order_idx']))

            self.summary_table_data = []
            self.generated_figs = []

            for item in sorted_display_data:
                fig = plt.Figure(figsize=(8, 4), dpi=100)
                ax = fig.add_subplot(111)
                ax.plot(df['Time'], df[item['plot_col_name']],
                        label=f'{item["param_name"]} ({item["sensor_num"]})')
                ax.set_title(f'{item["param_name"]} {self.t("plot_sensor")} {item["sensor_num"]}')
                ax.set_xlabel(self.t('plot_time_axis'))
                ax.set_ylabel(f'{item["param_name"]} ({item["unit"]})')
                ax.grid(True)
                ax.axhline(y=item['critical_threshold'], color='orange', linestyle='--',
                           label=f'{self.t("plot_critical_line")} '
                                 f'({item["critical_threshold"]:.2f} {item["unit"]})')
                ax.legend()

                chart_canvas = FigureCanvasTkAgg(fig, master=self.charts_frame)
                chart_canvas.draw()
                cw = chart_canvas.get_tk_widget()
                cw.pack(side=tk.TOP, fill=tk.BOTH, expand=True, pady=5)
                self.canvas_widgets.append(cw)
                self.generated_figs.append(fig)

                toolbar = NavigationToolbar2Tk(chart_canvas, self.charts_frame)
                toolbar.update()
                toolbar.pack(side=tk.TOP, fill=tk.X, pady=0)
                self.canvas_widgets.append(toolbar)

                if item['critical_time_text']:
                    lbl = tk.Label(self.charts_frame, text=item['critical_time_text'],
                                   fg="red", font=("Arial", 9))
                    lbl.pack(side=tk.TOP, pady=(0, 10))
                    self.canvas_widgets.append(lbl)

                self.summary_table_data.append({
                    "param":    item["param_name"],
                    "sensor":   item["sensor_num"],
                    "critical": item["critical_threshold"],
                    "time":     item["critical_time_for_summary"],
                })

            self.charts_frame.update_idletasks()
            self.canvas.config(scrollregion=self.canvas.bbox("all"))

            # Summary table widget
            if self.summary_table_data:
                summary_frame = tk.Frame(self.charts_frame, bd=2, relief="groove",
                                         padx=10, pady=10)
                summary_frame.pack(side=tk.TOP, fill=tk.X, pady=20)
                self.canvas_widgets.append(summary_frame)

                tk.Label(summary_frame, text=self.t('summary_title'),
                         font=("Arial", 12, "bold")).pack(pady=(0, 10))

                cols = (self.t('col_param'), self.t('col_sensor'),
                        self.t('col_critical'), self.t('col_time'))
                tree = ttk.Treeview(summary_frame, columns=cols, show="headings")
                for col in cols:
                    tree.heading(col, text=col)
                tree.column(cols[0], width=150, anchor="center")
                tree.column(cols[1], width=80, anchor="center")
                tree.column(cols[2], width=150, anchor="center")
                tree.column(cols[3], width=150, anchor="center")

                for item in self.summary_table_data:
                    tree.insert("", "end", values=(
                        item["param"], item["sensor"],
                        f"{item['critical']:.2f}", item["time"]))

                tree_vsb = ttk.Scrollbar(summary_frame, orient="vertical", command=tree.yview)
                tree_hsb = ttk.Scrollbar(summary_frame, orient="horizontal", command=tree.xview)
                tree.configure(yscrollcommand=tree_vsb.set, xscrollcommand=tree_hsb.set)
                tree_vsb.pack(side="right", fill="y")
                tree_hsb.pack(side="bottom", fill="x")
                tree.pack(expand=True, fill="both")
                for w in (tree, tree_vsb, tree_hsb):
                    self.canvas_widgets.append(w)

            self.status_label.config(text=self.t('status_done'), fg="green")
            messagebox.showinfo(self.t('success_title'), self.t('info_success'))
            self.btn_export_doc.config(state=tk.NORMAL)

        except FileNotFoundError as fnfe:
            self.status_label.config(text=self.t('status_no_data'), fg="red")
            messagebox.showerror(self.t('err_title'), self.t('err_file_not_found') + str(fnfe))
            self.btn_export_doc.config(state=tk.DISABLED)
        except Exception as e:
            self.status_label.config(text=self.t('err_unexpected') + str(e), fg="red")
            messagebox.showerror(self.t('err_title'), self.t('err_unexpected') + str(e))
            self.btn_export_doc.config(state=tk.DISABLED)

    # ── Word export ──────────────────────────────────────────────────────

    def export_results_to_doc(self):
        if not self.generated_figs or not self.summary_table_data:
            messagebox.showwarning(self.t('warn_title'), self.t('err_no_export'))
            return

        output_filepath = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Documents", "*.docx")],
            title=self.t('save_as_title'),
        )
        if not output_filepath:
            messagebox.showinfo(self.t('info_title'), self.t('info_cancelled_export'))
            return

        self.status_label.config(text=self.t('status_exporting'), fg="blue")
        self.master.update_idletasks()

        document = Document()
        document.add_heading(
            self.t('word_report_title') + os.path.basename(self.input_file_path), level=1)
        document.add_paragraph(
            self.t('word_created') + datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
        document.add_paragraph("")

        document.add_heading(self.t('word_charts_heading'), level=2)
        with tempfile.TemporaryDirectory() as tmpdir:
            for i, fig in enumerate(self.generated_figs):
                plot_filename = os.path.join(tmpdir, f"plot_{i}.png")
                fig.savefig(plot_filename, dpi=300, bbox_inches='tight')
                chart_title = (fig.axes[0].get_title()
                               if fig.axes and fig.axes[0].get_title()
                               else f"{self.t('word_chart_default')} {i + 1}")
                document.add_heading(chart_title, level=3)
                document.add_picture(plot_filename, width=Inches(6))
                document.add_paragraph("")

        document.add_heading(self.t('word_table_heading'), level=2)
        if self.summary_table_data:
            table = document.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = self.t('col_param')
            hdr[1].text = self.t('col_sensor')
            hdr[2].text = self.t('col_critical')
            hdr[3].text = self.t('col_time')
            for item in self.summary_table_data:
                row = table.add_row().cells
                row[0].text = str(item["param"])
                row[1].text = str(item["sensor"])
                row[2].text = f"{item['critical']:.2f}"
                row[3].text = str(item["time"])

        try:
            document.save(output_filepath)
            self.status_label.config(
                text=self.t('status_export_done') + os.path.basename(output_filepath),
                fg="green")
            messagebox.showinfo(self.t('success_title'),
                                self.t('info_export_done') + output_filepath)
        except Exception as e:
            self.status_label.config(text=self.t('err_export_word') + str(e), fg="red")
            messagebox.showerror(self.t('err_title'), self.t('err_export_word') + str(e))


if __name__ == "__main__":
    root = tk.Tk()
    app = FDSAnalyzerApp(root)
    root.mainloop()
