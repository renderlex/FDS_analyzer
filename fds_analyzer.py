import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os
from tkinter import Tk, filedialog # Для вибору файлів, хоч і без повноцінного GUI

def process_fds_data(file_path, output_dir="."):
    """
    Обробляє дані FDS, генерує графіки з критичними значеннями
    та експортує їх у файли Excel та Word.

    Args:
        file_path (str): Шлях до вхідного файлу FDS (.txt).
        output_dir (str): Директорія для збереження вихідних файлів.
    """
    try:
        # --- 1. Завантаження даних ---
        df = pd.read_csv(file_path, header=[0, 1], skipinitialspace=True)

        # --- 2. Парсинг заголовків та даних ---
        new_columns = []
        for col_header, sub_header in df.columns:
            if 'Unnamed' in sub_header:
                new_columns.append(col_header)
            else:
                sub_header = sub_header.replace('"', '')
                new_columns.append(f"{sub_header} ({col_header})")
        df.columns = new_columns
        df.rename(columns={'s': 'Time'}, inplace=True)

        # Перетворення всіх числових стовпців до числового типу
        for col in df.columns:
            if col != 'Time':
                df[col] = pd.to_numeric(df[col], errors='coerce')
        df.dropna(inplace=True) # Видаляємо рядки з NaN, якщо були помилки конвертації

        print("Дані успішно завантажені та оброблені.")
        print(f"Кількість рядків: {len(df)}")
        print(f"Кількість стовпців: {len(df.columns)}")
        print("\nПерші 5 рядків обробленого DataFrame:")
        print(df.head())
        print("\nНазви стовпців:")
        print(df.columns.tolist())

        # --- 3. Структурування даних та підготовка до графіків ---
        parameters_info = {
            'Temp': {'name': 'Температура', 'unit': 'C', 'critical': 60, 'is_upper_limit': True},
            'Visio': {'name': 'Видимість', 'unit': 'm', 'critical': 20, 'is_upper_limit': False},
            'TP': {'name': 'Тепловий потік', 'unit': 'kW/m2', 'critical': 2.5, 'is_upper_limit': True}, # 2500 Вт/м² = 2.5 кВт/м²
            'KK': {'name': 'Концентрація кисню', 'unit': 'kg/m3', 'critical': 0.232, 'is_upper_limit': False},
            'OV': {'name': 'Чадний газ (CO)', 'unit': 'kg/m3', 'critical': 1.16e-3, 'is_upper_limit': True},
            'DV': {'name': 'Двоокис вуглецю (CO2)', 'unit': 'kg/m3', 'critical': 0.11, 'is_upper_limit': True}
        }

        sensor_numbers = set()
        for col in df.columns:
            match = re.search(r'([A-Za-z]+)(\d+)', col)
            if match and match.group(1) in parameters_info:
                sensor_numbers.add(int(match.group(2)))

        max_sensor_number = max(sensor_numbers) if sensor_numbers else 0
        print(f"\nЗнайдено {max_sensor_number} датчиків.")
        if max_sensor_number == 0:
            print("Не знайдено стовпців з даними датчиків для обробки. Перевірте формат файлу.")
            return

        # --- 4. Побудова графіків та експорт в Word ---
        document = Document()
        document.add_heading('Результати моделювання пожежі FDS', level=1)

        # Додаємо інформацію про граничні значення
        critical_values_text = """
        Гранично допустимі значення за кожним з небезпечних чинників пожежі складають:
        — за підвищеною температурою — 60°C;
        — за тепловим потоком — 2.5 кВт/м² (2500 Вт/м²);
        — за втратою видимості — 20 м (у разі, коли обидва горизонтальні лінійні розміри приміщення
        менше ніж 20 м, гранично допустиму відстань щодо втрати видимості приймають рівною найбільшому
        горизонтальному лінійному розміру);
        — за зниженням вмісту кисню — 0.232 кг/м³;
        — за чадним газом (СО) — 1.16·10⁻³ кг/м³;
        — за двоокисом вуглецю (СО₂) — 0.11 кг/м³.
        """
        document.add_paragraph(critical_values_text)
        document.add_page_break()


        for sensor_num in range(1, max_sensor_number + 1):
            document.add_heading(f'Дані та графіки для Датчика {sensor_num}', level=2)

            for param_code, info in parameters_info.items():
                param_name = info['name']
                unit = info['unit']
                critical_value = info['critical']
                is_upper_limit = info['is_upper_limit']

                # Знаходимо точну назву стовпця з одиницями
                target_col = None
                for col in df.columns:
                    if col.startswith(f"{param_code}{sensor_num}") and f" ({unit})" in col:
                        target_col = col
                        break

                if target_col and target_col in df.columns:
                    plt.figure(figsize=(12, 7)) # Збільшуємо розмір для кращої видимості
                    plt.plot(df['Time'], df[target_col], label=f'{param_name} Датчика {sensor_num}', color='blue')

                    # Додаємо горизонтальну лінію критичного значення
                    plt.axhline(y=critical_value, color='red', linestyle='--', label=f'Критичне значення ({critical_value} {unit})')

                    # Позначаємо точки, де значення перевищує/опускається нижче критичного
                    if is_upper_limit:
                        critical_points = df[df[target_col] >= critical_value]
                    else: # Для кисню та видимості - значення нижче критичного
                        critical_points = df[df[target_col] <= critical_value]

                    if not critical_points.empty:
                        plt.scatter(critical_points['Time'], critical_points[target_col],
                                    color='red', zorder=5, label='Критичні точки')
                        
                        # Можна додати анотацію для першої критичної точки
                        first_critical_time = critical_points['Time'].iloc[0]
                        first_critical_value = critical_points[target_col].iloc[0]
                        
                        plt.annotate(f'Критичне значення досягнуто на {first_critical_time:.2f} с',
                                     xy=(first_critical_time, first_critical_value),
                                     xytext=(first_critical_time + 50, first_critical_value + (df[target_col].max() - df[target_col].min()) * 0.1), # Зміщення для анотації
                                     arrowprops=dict(facecolor='black', shrink=0.05),
                                     fontsize=9, color='red')


                    plt.xlabel('Час (с)', fontsize=12)
                    plt.ylabel(f"{param_name} ({unit})", fontsize=12)
                    plt.title(f"{param_name} для Датчика {sensor_num} (з критичними значеннями)", fontsize=14)
                    plt.grid(True, linestyle='--', alpha=0.7)
                    plt.legend()
                    plt.tight_layout() # Автоматично налаштовує розміри, щоб уникнути обрізки
                    
                    # Зберігаємо графік у тимчасовий файл
                    plot_filename = os.path.join(output_dir, f"plot_{param_code}_{sensor_num}.png")
                    plt.savefig(plot_filename)
                    plt.close()

                    # Додаємо графік у Word документ
                    p = document.add_paragraph()
                    r = p.add_run(f"Графік: {param_name} для Датчика {sensor_num}")
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    document.add_picture(plot_filename, width=Inches(6.5))
                    
                    # Додаємо текст, якщо були критичні точки
                    if not critical_points.empty:
                        min_time_critical = critical_points['Time'].min()
                        max_time_critical = critical_points['Time'].max()
                        p = document.add_paragraph(f"  * Критичне значення {param_name} досягнуто з {min_time_critical:.2f} с по {max_time_critical:.2f} с.")
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    document.add_page_break()

                else:
                    print(f"Попередження: Стовпець '{param_code}{sensor_num}' для параметра '{param_name}' не знайдено в даних.")

        # --- 5. Експорт в Excel (.xlsx) ---
        excel_output_path = os.path.join(output_dir, 'FDS_Results.xlsx')
        try:
            df.to_excel(excel_output_path, index=False)
            print(f"\nДані успішно експортовані в Excel: {excel_output_path}")
        except Exception as e:
            print(f"Помилка при експорті в Excel: {e}")

        # --- 6. Експорт в Word (.docx) ---
        word_output_path = os.path.join(output_dir, 'FDS_Analysis_Report.docx')
        try:
            document.save(word_output_path)
            print(f"Звіт з графіками успішно експортований в Word: {word_output_path}")
        except Exception as e:
            print(f"Помилка при експорті в Word: {e}")

        # Очищення тимчасових файлів графіків
        for sensor_num in range(1, max_sensor_number + 1):
            for param_code in parameters_info.keys():
                plot_filename = os.path.join(output_dir, f"plot_{param_code}_{sensor_num}.png")
                if os.path.exists(plot_filename):
                    os.remove(plot_filename)
        print("\nТимчасові файли графіків видалено.")

    except FileNotFoundError:
        print(f"Помилка: Файл '{file_path}' не знайдено. Перевірте шлях.")
    except Exception as e:
        print(f"Виникла помилка під час обробки даних: {e}")

# --- Виконання скрипта ---
if __name__ == "__main__":
    # Створюємо приховане вікно Tkinter для діалогу вибору файлу/папки
    root = Tk()
    root.withdraw() # Приховуємо головне вікно

    print("Будь ласка, оберіть файл FDS (.txt) для обробки.")
    input_file_path = filedialog.askopenfilename(
        title="Виберіть файл FDS (.txt)",
        filetypes=[("Text files", "*.txt")]
    )

    if input_file_path:
        print(f"Вибраний файл: {input_file_path}")
        print("Будь ласка, оберіть директорію для збереження результатів.")
        output_directory = filedialog.askdirectory(
            title="Виберіть директорію для збереження результатів"
        )

        if output_directory:
            print(f"Результати будуть збережені в: {output_directory}")
            process_fds_data(input_file_path, output_directory)
        else:
            print("Директорію для збереження не вибрано. Обробку скасовано.")
    else:
        print("Файл для обробки не вибрано. Обробку скасовано.")